import os
import re
import pandas as pd
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import requests

app = Flask(__name__)
CORS(app)

# --- CONFIGURATION ---
# REMEMBER: Switch to /current_routine.json before final deployment!
FIREBASE_URL = "https://edu-routine-generator-default-rtdb.asia-southeast1.firebasedatabase.app/current_routine.json"
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD")

# ==========================================
# --- HELPER FUNCTIONS ---
# ==========================================

def is_time_slot(text):
    """
    Returns True if the text looks like a time range (8.30-10.00).
    """
    clean_text = text.replace(' ', '').replace('\n', '')
    return bool(re.search(r'\d{1,2}[\.:]\d{2}[-–—]\d{1,2}[\.:]\d{2}', clean_text))

def get_faculty_mapping(doc):
    mapping = {}
    for table in doc.tables:
        rows = table.rows
        if not rows: continue
        
        for i in range(min(3, len(rows))):
            header = [cell.text.strip().lower() for cell in rows[i].cells]
            if 'name' in header and ('short form' in header or 'acronym' in header):
                try:
                    name_idx = header.index('name')
                    acronym_idx = -1
                    for idx, h in enumerate(header):
                        if 'short' in h or 'acronym' in h:
                            acronym_idx = idx
                            break
                    
                    if acronym_idx != -1:
                        for row in rows[i+1:]:
                            cells = row.cells
                            if len(cells) > max(name_idx, acronym_idx):
                                name = cells[name_idx].text.strip()
                                acronym = cells[acronym_idx].text.strip()
                                if acronym and name:
                                    mapping[acronym] = name
                except:
                    continue
                break
    return mapping

def parse_routine_complete(doc):
    data = []
    faculty_map = get_faculty_mapping(doc)
    
    # Regex Patterns
    course_pattern = re.compile(r"([A-Z]{2,4}\s*\d{3}(?:\.[0-9A-Za-z]+)?)", re.IGNORECASE)
    # New: Regex to capture specific times inside cells (e.g. 1.30-3.30 or 1:30 - 3:30)
    # Capture group 1 is the time string
    custom_time_pattern = re.compile(r"\(?(\d{1,2}[\.:]\d{2}\s*[-–—]\s*\d{1,2}[\.:]\d{2})\)?")

    valid_days = ['Saturday', 'Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    current_day = None 

    for table in doc.tables:
        rows = table.rows
        if not rows: continue
        
        # Check Day at top of table
        top_cell = rows[0].cells[0].text.strip()
        if top_cell in valid_days:
            current_day = top_cell
        
        # Find Header Row
        header_row_index = -1
        time_slots = []
        
        for i in range(min(5, len(rows))):
            cells = [cell.text.strip() for cell in rows[i].cells]
            found_slots = []
            for idx, text in enumerate(cells):
                if is_time_slot(text):
                    found_slots.append({'index': idx, 'time': text.strip()})
            
            if len(found_slots) >= 1:
                header_row_index = i
                time_slots = found_slots
                break
        
        if header_row_index == -1:
            continue

        # Parse Rows
        for row in rows[header_row_index+1:]:
            cells = row.cells
            if not cells: continue
            
            first_cell_text = cells[0].text.strip()
            
            if first_cell_text in valid_days:
                current_day = first_cell_text
                continue
            
            if not current_day:
                continue

            # Check for Room Number
            if first_cell_text.isdigit() or first_cell_text.startswith('N'):
                room = first_cell_text
                
                for slot in time_slots:
                    if slot['index'] < len(cells):
                        raw_text = cells[slot['index']].text.strip()
                        if not raw_text: continue
                        
                        # Use a working copy of text for cleaning
                        processing_text = raw_text.replace('\n', ' ')
                        
                        # --- 1. DETECT CUSTOM TIME OVERRIDE ---
                        final_time = slot['time']
                        
                        # Check if cell contains a specific time (e.g., "(1.30-3.30)")
                        time_match = custom_time_pattern.search(processing_text)
                        if time_match:
                            # Use the found time (e.g. "1.30-3.30")
                            final_time = time_match.group(1).replace(' ', '') # Normalize spaces
                            # Remove the time string (and surrounding parens) from text so it's not treated as Faculty
                            # We replace the FULL match (including parens if regex caught them)
                            processing_text = processing_text.replace(time_match.group(0), '')

                        # --- 2. EXTRACT COURSE ---
                        course_match = course_pattern.search(processing_text)
                        
                        if course_match:
                            course_code = course_match.group(1).strip()
                            
                            # --- 3. EXTRACT FACULTY (CLEANING STRATEGY) ---
                            # Remove Course Code
                            remaining = processing_text.replace(course_code, '')
                            
                            # Remove "(EEE)" or "(CSE)" noise
                            remaining = re.sub(r'\([A-Z]{3}\)', '', remaining)
                            
                            # Remove empty parentheses "()"
                            remaining = remaining.replace('()', '')
                            
                            # Clean up commas, spaces, dashes
                            faculty_acronym = remaining.strip(' ,-–—')
                            
                            # Final sanity check: Faculty shouldn't be too long
                            if len(faculty_acronym) > 15:
                                faculty_acronym = "" # Logic failed, safer to show nothing than garbage

                            # Determine Type
                            header_text_full = " ".join([c.text.lower() for c in rows[header_row_index].cells])
                            prev_row_text = rows[max(0, header_row_index-1)].cells[0].text.lower() if header_row_index > 0 else ""
                            
                            class_type = 'Theory'
                            if 'lab' in header_text_full or 'lab' in prev_row_text:
                                class_type = 'Lab'

                            data.append({
                                "Day": current_day,
                                "Time": final_time, # Uses override if found
                                "Room": room,
                                "Course": course_code,
                                "FacultyAcronym": faculty_acronym,
                                "FacultyFullName": faculty_map.get(faculty_acronym, ""),
                                "Type": class_type
                            })
    
    if not data:
        return pd.DataFrame(columns=["Day", "Time", "Room", "Course", "FacultyAcronym", "FacultyFullName", "Type"])
        
    return pd.DataFrame(data)

# ==========================================
# --- ROUTES ---
# ==========================================

@app.route('/api/upload', methods=['POST'])
def upload_file():
    try:
        user_password = request.form.get('password')
        if ADMIN_PASSWORD and user_password != ADMIN_PASSWORD:
            return jsonify({"error": "Wrong password"}), 403

        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        doc = Document(file.stream)
        df = parse_routine_complete(doc)
        
        routine_json = df.to_dict(orient="records")
        
        response = requests.put(FIREBASE_URL, json={
            "updatedAt": str(pd.Timestamp.now()),
            "data": routine_json
        })

        return jsonify({
            "status": "success", 
            "count": len(routine_json),
            "firebase_status": response.status_code
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(port=8000, debug=True)